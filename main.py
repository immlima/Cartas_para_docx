import os
from tkinter import filedialog
from PIL import Image
from docx import Document
from docx.shared import Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

document = Document()
sections = document.sections
for section in sections:
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(0)
    section.right_margin = Cm(0)


linhabranca=1
name_3x3=0
i=0
j=0
k=1
dimecaow=3
dimecaoh=3

pixelW=748
pixelH=1048
img3x3 = Image.new ('RGB',(dimecaow*(pixelW+linhabranca),dimecaoh*(pixelH+linhabranca)), (255,255,255))


pasta= filedialog.askdirectory(title=f"Escolha o diretorio onde estam as imagens das cartas")
# print(pasta)
pasta3x3=pasta+" 3x3"
if not os.path.exists(pasta3x3):
    os.mkdir(pasta3x3)
# print(pasta3x3)
#pasta3x3=filedialog.askdirectory(title=f"Escolha um diretorio para salvar as imagens {dimecaow}x{dimecaoh} e o .Docx")


for root, dirs, files in os.walk(pasta):
    for file in files:
        original_img_path = os.path.join(root, file)
        pillow_img = Image.open(original_img_path)
        width, height = pillow_img.size
        if width!=pixelW or height!=pixelH:
            new_img = pillow_img.resize((pixelW,pixelH), Image.LANCZOS)
            new_img.save(original_img_path,optimize=True,quality=100)
            
        
        if i>dimecaow-1:
            j =j+1
            i=0
        if j>dimecaoh-1:
            i=0
            j=0

        pillow_img = Image.open(original_img_path)
        img3x3.paste (pillow_img, (i*(pixelW+linhabranca),j*(pixelH+linhabranca)))

        #print(k, len(files))
        if (i==dimecaow-1 and j==dimecaoh-1) or k==len(files):

            print("Salvou com: "+str(k)+' '+str(len(files)) )

            name_3x3_str=str(name_3x3)
            #if len(str(name_3x3))==1:
            #    name_3x3_str="00"+str(name_3x3)
            #else:
            #    name_3x3_str="0"+str(name_3x3)
            img3x3.save(os.path.join(pasta3x3, 'Cartas '+name_3x3_str+'.png'),optimize=True,quality=100)
            img3x3 = Image.new ('RGB',(dimecaow*pixelW,dimecaoh*pixelH), (255,255,255))

            name_3x3=name_3x3+1
        i =i+1
        k =k+1

files_sort=[]
for root, dirs, files in os.walk(pasta3x3, topdown=True):
    
    files_sort=files
    def myFunc(e):
        return len(e)


    files_sort.sort(key=myFunc)
    
    #print(files_sort)
    for file in files_sort:
        #print(file)
        
        excesion=('docx', 'pdf')
        if file.split(".")[-1] in excesion:
            continue   
        original_img_path = os.path.join(root, file)
        #print(original_img_path)
        document.add_picture(original_img_path, width=Inches(dimecaow*2.5), height=Inches(dimecaoh*3.5))
        last_paragraph = document.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        #print(f"nova imagem: {original_img_path}")

document.save(os.path.join(root, pasta.split("/")[-1]+'.docx'))
print('Arquivo .docx salvo')



'''
for root, dirs, files in os.walk("C:\\Users\\Mateus\\Downloads\\RPG\\DnD\\5E\Monstros\\[D&D 5E] - Cartas - Bestiario dos Monstros - Ordenar Paginas"):
    for file in files:
        os.rename(os.path.join(root,file),os.path.join(root,'Bestiario dos Monstros '+file[-8:]))
'''   

#Exclui as imagens 3x3
"""
for root, dirs, files in os.walk(pasta3x3):
    for file in files:
        
        original_img_path = os.path.join(root, file)
        excesion=('png', 'jpeg', 'jpg')
        if file.split(".")[-1] in excesion:
            if os.path.exists(original_img_path):
                os.remove(original_img_path)
 """              


"""
for root, dirs, files in os.walk(pasta):
    for file in files:
        original_img_path = os.path.join(root, file)
        #print(original_img_path)
        pillow_img = Image.open(original_img_path)
        width, height = pillow_img.size
        if width!=pixelW or height!=pixelH:
            new_img = pillow_img.resize((pixelW,pixelH), Image.LANCZOS)
            new_img.save(original_img_path,optimize=True,quality=100)
"""